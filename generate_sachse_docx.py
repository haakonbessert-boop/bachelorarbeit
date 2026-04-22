#!/usr/bin/env python3
"""Generate sachse_gliederung.docx with clean Word-native formatting."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(3.0)
section.bottom_margin = Cm(2.0)

# ── Default font ────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)


# ── Helpers ─────────────────────────────────────────────────────────────────
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_horizontal_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr %s>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        '</w:pBdr>' % nsdecls('w')
    )
    pPr.append(pBdr)


def set_cell_text(cell, text, bold=False, size=None, alignment=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = size or Pt(10)
    run.bold = bold


def shade_row(row, color="D9E2F3"):
    for cell in row.cells:
        shading = parse_xml(
            '<w:shd %s w:fill="%s"/>' % (nsdecls('w'), color)
        )
        cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml('<w:tblPr %s/>' % nsdecls('w'))
    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="808080"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)


# ── Titelblock ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Gliederung und Zeitplan")
run.bold = True
run.font.size = Pt(16)
run.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Stand: 22. April 2026")
run.font.size = Pt(10)
run.font.name = "Calibri"

add_horizontal_line()

meta_data = [
    ("Titel", "Ein Mehrkriterien-Entscheidungsmodell zur Auswahl eines KPI-Hub-Integrationsansatzes in einer fragmentierten Enterprise-Systemlandschaft"),
    ("Verfasser", "Haakon Bessert (Matrikelnr. 3005149)"),
    ("Praxispartner", "SAP SE — Janine Steidelmüller"),
    ("Betreuer DHSN", "Prof. Dr.-Ing. Jürgen Sachse"),
    ("Abgabe", "10. Juli 2026"),
]

meta_table = doc.add_table(rows=len(meta_data), cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (label, value) in enumerate(meta_data):
    set_cell_text(meta_table.rows[i].cells[0], label, bold=True)
    set_cell_text(meta_table.rows[i].cells[1], value)
meta_table.columns[0].width = Cm(3.5)
meta_table.columns[1].width = Cm(12.5)

add_horizontal_line()

# ── Kapitelübersicht ────────────────────────────────────────────────────────
add_heading_styled("Gliederung", level=1)
doc.add_paragraph("Die Arbeit umfasst 10 Kapitel auf 62 Seiten. Der detaillierte Aufbau ist nachfolgend dargestellt.")

chapters = [
    ("1", "Einleitung", "4"),
    ("2", "Grundlagen und Stand der Forschung", "9"),
    ("3", "Methodik", "6"),
    ("4", "Ausgangssituation und Systemlandschaft", "7"),
    ("5", "Anforderungsanalyse", "7"),
    ("6", "Lösungsraum und Architektur-/Tooloptionen", "6"),
    ("7", "Entscheidungs- und Bewertungsmodell", "8"),
    ("8", "MVP: Design, Umsetzung und Validierung", "8"),
    ("9", "Ergebnisse und Diskussion", "5"),
    ("10", "Fazit und Ausblick", "2"),
]

ch_table = doc.add_table(rows=len(chapters) + 2, cols=3)
set_table_borders(ch_table)
ch_table.alignment = WD_TABLE_ALIGNMENT.LEFT

header = ch_table.rows[0]
set_cell_text(header.cells[0], "Nr.", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
set_cell_text(header.cells[1], "Kapitel", bold=True)
set_cell_text(header.cells[2], "S.", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
shade_row(header)

for i, (nr, name, pages) in enumerate(chapters):
    row = ch_table.rows[i + 1]
    set_cell_text(row.cells[0], nr, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_text(row.cells[1], name)
    set_cell_text(row.cells[2], pages, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

total_row = ch_table.rows[-1]
set_cell_text(total_row.cells[0], "")
set_cell_text(total_row.cells[1], "Gesamt", bold=True)
set_cell_text(total_row.cells[2], "62", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
shade_row(total_row)

ch_table.columns[0].width = Cm(1.5)
ch_table.columns[1].width = Cm(12.0)
ch_table.columns[2].width = Cm(1.5)

# ── Detaillierte Untergliederung ────────────────────────────────────────────
add_heading_styled("Detaillierte Untergliederung", level=2)

left_col = [
    (True,  "1  Einleitung"),
    (False, "1.1  Ausgangssituation und Problemstellung"),
    (False, "1.2  Lösungsrelevanz und Motivation"),
    (False, "1.3  Einordnung in die Wirtschaftsinformatik"),
    (False, "1.4  Zielsetzung und Forschungsfrage"),
    (False, "1.5  Aufbau der Arbeit"),
    (False, ""),
    (True,  "2  Grundlagen und Stand der Forschung"),
    (False, "2.1  KPI-Management und Business Intelligence"),
    (False, "2.2  Enterprise-Systemintegration"),
    (False, "2.3  Data Governance und Zugriffskontrolle"),
    (False, "2.4  Entscheidungsmethoden im Software Engineering"),
    (False, "2.5  Defizite bestehender Ansätze"),
    (False, ""),
    (True,  "3  Methodik"),
    (False, "3.1  Forschungsansatz (Design Science Research)"),
    (False, "3.2  Literaturrecherche"),
    (False, "3.3  Ergebniserwartung"),
    (False, "3.4  Vorgehen im Überblick"),
    (False, ""),
    (True,  "4  Ausgangssituation und Systemlandschaft"),
    (False, "4.1  Organisatorischer Kontext"),
    (False, "4.2  Bestehende Systemlandschaft"),
    (False, "4.3  Ist-Zustand der KPI-Versorgung"),
    (False, "4.4  Datenbereitstellung und Schnittstellenlandschaft"),
    (False, "4.5  Identifizierte Problemfelder"),
    (False, ""),
    (True,  "5  Anforderungsanalyse"),
    (False, "5.1  Vorgehen bei der Anforderungserhebung"),
    (False, "5.2  Stakeholder und Use Cases"),
    (False, "5.3  Funktionale Anforderungen"),
    (False, "5.4  Nicht-funktionale Anforderungen"),
    (False, "5.5  Priorisierung der Anforderungen"),
]

right_col = [
    (True,  "6  Lösungsraum und Architektur-/Tooloptionen"),
    (False, "6.1  Identifikation relevanter Lösungsansätze"),
    (False, "6.2  Portal-/Dashboard-Föderation (Embedding)"),
    (False, "6.3  Zentrale BI-Aggregation"),
    (False, "6.4  Low-Code / No-Code Plattformen"),
    (False, "6.5  Kundenspezifische Webanwendung"),
    (False, "6.6  Referenzprojekte bei SAP"),
    (False, "6.7  Vergleichende Übersicht und Machbarkeitseinschätzung"),
    (False, ""),
    (True,  "7  Entscheidungs- und Bewertungsmodell"),
    (False, "7.1  Wahl der Bewertungsmethodik"),
    (False, "7.2  Aufbau des Scoring-Modells"),
    (False, "7.3  Bewertung der Lösungsoptionen"),
    (False, "7.4  Ergebnis und Empfehlung"),
    (False, "7.5  Sensitivitätsanalyse"),
    (False, ""),
    (True,  "8  MVP: Design, Umsetzung und Validierung"),
    (False, "8.1  MVP-Zielsetzung und Scope"),
    (False, "8.2  Architektur und Design"),
    (False, "8.3  Umsetzung"),
    (False, "8.4  Validierung"),
    (False, ""),
    (True,  "9  Ergebnisse und Diskussion"),
    (False, "9.1  Ergebnisse der Bewertung"),
    (False, "9.2  Beantwortung der Forschungsfrage"),
    (False, "9.3  Kritische Diskussion"),
    (False, "9.4  Vergleich mit Ergebniserwartung"),
    (False, ""),
    (True,  "10  Fazit und Ausblick"),
    (False, "10.1  Fazit"),
    (False, "10.2  Beitrag der Arbeit"),
    (False, "10.3  Ausblick"),
    (False, ""),
    (False, ""),
]

num_rows = max(len(left_col), len(right_col))
while len(left_col) < num_rows:
    left_col.append((False, ""))
while len(right_col) < num_rows:
    right_col.append((False, ""))

detail_table = doc.add_table(rows=num_rows, cols=2)
detail_table.alignment = WD_TABLE_ALIGNMENT.LEFT

for i in range(num_rows):
    l_bold, l_text = left_col[i]
    r_bold, r_text = right_col[i]
    set_cell_text(detail_table.rows[i].cells[0], l_text, bold=l_bold, size=Pt(9) if not l_bold else Pt(10))
    set_cell_text(detail_table.rows[i].cells[1], r_text, bold=r_bold, size=Pt(9) if not r_bold else Pt(10))

detail_table.columns[0].width = Cm(8.0)
detail_table.columns[1].width = Cm(8.0)

add_horizontal_line()

# ── Roter Faden ─────────────────────────────────────────────────────────────
add_heading_styled("Roter Faden", level=1)

p = doc.add_paragraph()
run = p.add_run("Die Kapitelstruktur folgt dem DSR-Prozessmodell nach Peffers et al. (2007) — von der Problemdefinition über die Artefaktkonstruktion bis zur Evaluation:")
run.font.size = Pt(11)
run.font.name = "Calibri"

doc.add_paragraph()

flow_table = doc.add_table(rows=2, cols=9)
flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

boxes = [
    "Kap. 1–4\nProblem &\nSystemlandschaft",
    "→",
    "Kap. 5\nAnforderungs-\nanalyse",
    "→",
    "Kap. 6–7\nLösungsraum &\nBewertung",
    "→",
    "Kap. 8–9\nMVP &\nValidierung",
    "→",
    "Kap. 10\nFazit &\nAusblick",
]

for i, text in enumerate(boxes):
    cell = flow_table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if text == "→":
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.bold = True
    else:
        for j, line in enumerate(text.split("\n")):
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.bold = (j == 0)
            if j < len(text.split("\n")) - 1:
                p.add_run("\n").font.size = Pt(9)
        shading = parse_xml('<w:shd %s w:fill="D9E2F3"/>' % nsdecls('w'))
        cell._tc.get_or_add_tcPr().append(shading)

phases = [
    ("Relevance Cycle", False),
    ("", False),
    ("", False),
    ("", False),
    ("Rigor Cycle", False),
    ("", False),
    ("", False),
    ("", False),
    ("", False),
]

# Merge cells for DSR phases row
phase_row = flow_table.rows[1]
# Relevance Cycle: col 0
set_cell_text(phase_row.cells[0], "Relevance Cycle", size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.CENTER)
# Rigor Cycle: cols 2-4
phase_row.cells[2].merge(phase_row.cells[4])
set_cell_text(phase_row.cells[2], "Rigor Cycle", size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.CENTER)
# Artefakt-Evaluation: col 6
set_cell_text(phase_row.cells[6], "Artefakt-Evaluation", size=Pt(8), alignment=WD_ALIGN_PARAGRAPH.CENTER)

add_horizontal_line()

# ── Strukturelle Entscheidungen ─────────────────────────────────────────────
add_heading_styled("Strukturelle Entscheidungen", level=1)

decisions = [
    ("DSR als Methodik", "Zwei Artefakte im Zentrum: Scoring-Modell + MVP"),
    ("Kap. 2.5 Entscheidungsmethoden\nim Software Engineering", "Theoretische Basis für Kap. 7 — Methodenwahl (AHP, Nutzwertanalyse etc.) wird hier fundiert"),
    ("Kap. 4.4 Datenbereitstellung", "Janine: explizit abzudeckendes Thema"),
    ("Kap. 6.6/6.7 Referenzprojekte +\nMachbarkeitseinschätzung", "Janines Kernkriterium: Umsetzbarkeit als Vorfilter vor dem Scoring"),
    ("Kap. 7.1 Wahl der\nBewertungsmethodik", "Sachse: Warum dieses Bewertungsverfahren? — explizit begründen"),
]

dec_table = doc.add_table(rows=len(decisions) + 1, cols=2)
set_table_borders(dec_table)
dec_table.alignment = WD_TABLE_ALIGNMENT.LEFT

header = dec_table.rows[0]
set_cell_text(header.cells[0], "Entscheidung", bold=True)
set_cell_text(header.cells[1], "Begründung", bold=True)
shade_row(header)

for i, (decision, reason) in enumerate(decisions):
    row = dec_table.rows[i + 1]
    set_cell_text(row.cells[0], decision, bold=True, size=Pt(10))
    set_cell_text(row.cells[1], reason, size=Pt(10))

dec_table.columns[0].width = Cm(5.5)
dec_table.columns[1].width = Cm(10.5)

add_horizontal_line()

# ── Zeitplan ────────────────────────────────────────────────────────────────
add_heading_styled("Zeitplan", level=1)

p = doc.add_paragraph()
run = p.add_run("Das Schreiben der Kapitel läuft parallel zu jeder Phase — nicht erst am Ende.")
run.font.size = Pt(11)
run.font.name = "Calibri"
run.italic = True

schedule = [
    ("Orientierung", "20. Apr – 7. Mai", "Onboarding, Systemlandschaft, Stakeholder", "Kap. 1, 3 + Literatur"),
    ("Anforderungen", "8. Mai – 24. Mai", "Interviews, Requirements erheben", "Kap. 2, 4"),
    ("Analyse & Modell", "25. Mai – 14. Jun", "Lösungsraum, Scoring-Modell", "Kap. 5, 6, 7"),
    ("MVP", "15. Jun – 30. Jun", "MVP umsetzen und validieren", "Kap. 8, 9, 10"),
    ("Puffer", "1. Jul – 9. Jul", "Korrekturen, Feinschliff", "—"),
    ("Abgabe", "10. Juli 2026", "", ""),
]

sch_table = doc.add_table(rows=len(schedule) + 1, cols=4)
set_table_borders(sch_table)
sch_table.alignment = WD_TABLE_ALIGNMENT.LEFT

header = sch_table.rows[0]
set_cell_text(header.cells[0], "Phase", bold=True)
set_cell_text(header.cells[1], "Zeitraum", bold=True)
set_cell_text(header.cells[2], "Schwerpunkt", bold=True)
set_cell_text(header.cells[3], "Parallel schreiben", bold=True)
shade_row(header)

for i, (phase, period, focus, writing) in enumerate(schedule):
    row = sch_table.rows[i + 1]
    is_deadline = (phase == "Abgabe")
    set_cell_text(row.cells[0], phase, bold=is_deadline, size=Pt(10))
    set_cell_text(row.cells[1], period, bold=is_deadline, size=Pt(10))
    set_cell_text(row.cells[2], focus, size=Pt(10))
    set_cell_text(row.cells[3], writing, size=Pt(10))
    if is_deadline:
        shade_row(row, "FFF2CC")

sch_table.columns[0].width = Cm(3.0)
sch_table.columns[1].width = Cm(3.5)
sch_table.columns[2].width = Cm(5.5)
sch_table.columns[3].width = Cm(4.0)

# ── Save ────────────────────────────────────────────────────────────────────
doc.save("sachse_gliederung.docx")
print("sachse_gliederung.docx generated successfully.")
